from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "docs" / "dossier_technique.md"
OUTPUT = ROOT / "docs" / "dossier_technique.docx"
ACCENT = "168C91"
DARK = "173F43"
LIGHT = "E8F3F3"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction_node, separate, end))


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^]]+\]\([^)]+\))")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(DARK)
        else:
            link_match = re.match(r"\[([^]]+)\]\(([^)]+)\)", token)
            if link_match:
                paragraph.add_run(f"{link_match.group(1)} ({link_match.group(2)})")
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    heading_sizes = {1: 18, 2: 14, 3: 11.5}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DARK if level == 1 else ACCENT)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.right_indent = Cm(0.5)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("AgriVista Field — Dossier technique  •  ")
    add_field(footer, "PAGE")

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    properties = document.core_properties
    properties.title = "AgriVista Field — Dossier technique"
    properties.subject = "Documentation technique de l’application Flutter AgriVista Field"
    properties.author = "Dorian Argaillot"
    properties.keywords = "Flutter, Dart, Riverpod, Dio, Hive, Clean Architecture"


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Cm(5.2)
    run = paragraph.add_run("AgriVista Field")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    run = subtitle.add_run("DOSSIER TECHNIQUE")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Cm(2.2)
    metadata.add_run("Développement mobile Flutter / Dart 3\n").bold = True
    metadata.add_run("Master BIHAR — Session 2025-2026\n")
    metadata.add_run("Version 1.0 — 4 septembre 2026")

    document.add_page_break()
    toc_title = document.add_heading("Table des matières", level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph("Dans Word : clic droit sur la table puis « Mettre à jour les champs ». ")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    document.add_page_break()


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F6F6")
    properties.append(shading)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            add_inline(cell.paragraphs[0], value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_index == 0:
                set_cell_shading(cell, ACCENT)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def convert_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            index += 1
            continue
        if stripped.startswith("|" ):
            raw_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    raw_rows.append(cells)
                index += 1
            add_table(document, raw_rows)
            continue
        if re.match(r"^- ", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        ordered = re.match(r"^\d+\.\s+(.+)", stripped)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, ordered.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith("```")
                or candidate.startswith("|")
                or candidate.startswith("- ")
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    configure_document(document)
    add_cover(document)
    convert_markdown(document, markdown)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
